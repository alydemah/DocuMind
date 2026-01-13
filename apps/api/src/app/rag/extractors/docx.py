from pathlib import Path

from docx import Document as DocxDocument

from app.rag.extractors.base import BaseExtractor, ExtractedPage, ExtractionResult


class DocxExtractor(BaseExtractor):
    def extract(self, file_path: Path) -> ExtractionResult:
        doc = DocxDocument(str(file_path))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        cleaned = self._clean_text(full_text)

        pages = [
            ExtractedPage(
                page_number=1,
                content=cleaned,
                metadata={"format": "docx"},
            )
        ]

        return ExtractionResult(
            pages=pages,
            total_pages=1,
            metadata={"format": "docx", "paragraph_count": len(paragraphs)},
        )
